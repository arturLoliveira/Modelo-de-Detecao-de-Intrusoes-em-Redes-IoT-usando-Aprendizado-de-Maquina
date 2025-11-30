from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Criar documento
doc = Document()

# Configurar margens (ABNT: 3cm esquerda, 2cm direita, 3cm superior, 2cm inferior)
sections = doc.sections
for section in sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# Função para adicionar título
def add_title(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(14)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(12)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.bold = True
        p.space_before = Pt(12)
        p.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        p.space_before = Pt(6)
        p.space_after = Pt(6)
    return p

# Função para adicionar parágrafo
def add_paragraph(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    
    # Processar formatação (negrito e itálico)
    parts = text.split('<b>')
    for i, part in enumerate(parts):
        if '</b>' in part:
            bold_text, rest = part.split('</b>', 1)
            run = p.add_run(bold_text)
            run.font.bold = True
            run.font.size = Pt(12)
            
            # Processar itálico no resto
            if '<i>' in rest:
                process_italic(p, rest)
            else:
                run = p.add_run(rest)
                run.font.size = Pt(12)
        else:
            if '<i>' in part:
                process_italic(p, part)
            else:
                run = p.add_run(part)
                run.font.size = Pt(12)
    return p

def process_italic(paragraph, text):
    parts = text.split('<i>')
    for j, ipart in enumerate(parts):
        if '</i>' in ipart:
            italic_text, rest = ipart.split('</i>', 1)
            run = paragraph.add_run(italic_text)
            run.font.italic = True
            run.font.size = Pt(12)
            if rest:
                run = paragraph.add_run(rest)
                run.font.size = Pt(12)
        else:
            run = paragraph.add_run(ipart)
            run.font.size = Pt(12)

# Adicionar bullet point
def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    process_italic(p, text)

# Adicionar numbered list
def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    process_italic(p, text)

# CONTEÚDO

# Título principal
add_title("3. MATERIAIS E MÉTODOS", level=1)

# 3.1
add_title("3.1 Natureza da Pesquisa", level=2)
add_paragraph("Esta pesquisa caracteriza-se como um estudo quantitativo e experimental, fundamentado na aplicação de algoritmos de aprendizado de máquina para detecção de intrusões em redes IoT. A abordagem quantitativa justifica-se pela necessidade de avaliar numericamente o desempenho dos modelos através de métricas estatísticas consolidadas na literatura, como acurácia, precisão, <i>recall</i> e F1-<i>score</i>. O caráter experimental manifesta-se no desenvolvimento controlado de modelos computacionais, na manipulação sistemática de hiperparâmetros e na análise comparativa de resultados obtidos em condições padronizadas e reprodutíveis.")

# 3.2
add_title("3.2 Datasets Utilizados", level=2)
add_paragraph("Para treinamento e validação dos modelos propostos, foram utilizados dois <i>datasets</i> públicos amplamente reconhecidos pela comunidade científica internacional:")

add_paragraph("<b>ToN-IoT (<i>Telemetry dataset of IoT</i>)</b>: Desenvolvido pela <i>University of New South Wales</i> (UNSW), Austrália, especificamente projetado para pesquisas em segurança cibernética de redes IoT e IIoT. Neste trabalho, utilizou-se o subconjunto <i>Network Dataset</i>, contendo 461.043 registros de tráfego de rede com 43 características cada. Os dados estão classificados em tráfego benigno e 9 categorias de ataques contemporâneos: <i>Denial of Service</i> (DoS), <i>Distributed Denial of Service</i> (DDoS), <i>ransomware</i>, <i>backdoor</i>, <i>injection attacks</i>, <i>scanning</i>, <i>cross-site scripting</i> (XSS), <i>password cracking</i> e <i>man-in-the-middle</i> (MitM).")

add_paragraph("<b>UNSW-NB15</b>: <i>Dataset</i> desenvolvido pela mesma instituição, contendo 257.673 registros de tráfego de rede divididos em conjunto de treinamento pré-definido (175.341 amostras) e conjunto de teste (82.332 amostras). Inclui tráfego normal e 9 categorias de ataques modernos: <i>Fuzzers</i>, <i>Analysis</i>, <i>Backdoors</i>, DoS, <i>Exploits</i>, <i>Generic</i>, <i>Reconnaissance</i>, <i>Shellcode</i> e <i>Worms</i>. Cada registro possui 49 características extraídas do tráfego de rede.")

add_paragraph("Para este trabalho, ambos os <i>datasets</i> foram tratados como problemas de <b>classificação binária</b> (Normal <i>vs.</i> Ataque), consolidando todas as categorias de ataques em uma única classe maliciosa. Esta abordagem permite comparação direta entre algoritmos e simplifica a implementação em cenários reais onde a distinção entre tipos específicos de ataque é menos crítica que a detecção da anomalia em si.")

# 3.3
add_title("3.3 Ambiente Computacional e Ferramentas", level=2)
add_paragraph("Todo o desenvolvimento, implementação e experimentação foram realizados utilizando a linguagem de programação <b>Python 3.9</b>. As principais bibliotecas e suas respectivas versões foram:")

add_bullet("<b>Pandas 1.5.3</b>: Manipulação e análise de dados tabulares")
add_bullet("<b>NumPy 1.23.5</b>: Operações matemáticas e <i>arrays</i> multidimensionais")
add_bullet("<b>Scikit-learn 1.2.1</b>: Pré-processamento, SVM e métricas de avaliação")
add_bullet("<b>XGBoost 1.7.3</b>: Implementação otimizada do algoritmo XGBoost")
add_bullet("<b>LightGBM 3.3.5</b>: Implementação eficiente do LightGBM")
add_bullet("<b>Matplotlib 3.6.3</b> e <b>Seaborn 0.12.2</b>: Visualização de dados e geração de gráficos")

add_paragraph("Os experimentos foram executados em ambiente local com sistema operacional Windows 11, processador Intel Core i7 (12ª geração) e 16GB de memória RAM. Todos os <i>scripts</i> foram organizados em arquivos Python independentes, cada um responsável por treinar um modelo específico em um <i>dataset</i> determinado.")

# 3.4
add_title("3.4 Pré-processamento de Dados", level=2)
add_paragraph("O <i>pipeline</i> de pré-processamento foi estruturado para garantir qualidade dos dados e evitar vazamento de informação (<i>data leakage</i>), seguindo as etapas:")

# 3.4.1
add_title("3.4.1 Remoção de Colunas Sensíveis", level=3)
add_paragraph("Identificadores que poderiam causar memorização espúria ou vazamento de informação foram sistematicamente removidos: identificadores temporais (ts, timestamp, date, time), endereços de rede (src_ip, dst_ip, src_port, dst_port), identificadores únicos (id) e, no caso do ToN-IoT, a coluna <i>type</i> que continha explicitamente o nome do tipo de ataque. Esta etapa é crítica para garantir que os modelos aprendam padrões comportamentais genuínos do tráfego de rede.")

# 3.4.2
add_title("3.4.2 Tratamento de Variáveis Categóricas", level=3)
add_paragraph("Variáveis categóricas foram convertidas para formato numérico via <b>Label Encoding</b>, com concatenação de todos os valores únicos presentes nos conjuntos de treino, validação e teste, tratamento de valores ausentes e aplicação da transformação de forma consistente.")

# 3.4.3
add_title("3.4.3 Imputação e Normalização", level=3)
add_paragraph("Para variáveis numéricas com valores faltantes, utilizou-se o <b>SimpleImputer</b> com estratégia pela média. A normalização foi realizada através do <b>StandardScaler</b>, transformando todas as características para média zero e desvio padrão unitário, etapa crítica especialmente para o algoritmo SVM. O <i>scaler</i> foi ajustado apenas no conjunto de treinamento e aplicado aos demais conjuntos.")

# 3.4.4
add_title("3.4.4 Divisão dos Dados", level=3)
add_paragraph("<b>ToN-IoT</b>: Divisão estratificada em 70% treinamento (323.330 amostras), 15% validação (69.157 amostras) e 15% teste (69.156 amostras). <b>UNSW-NB15</b>: Utilizados conjuntos pré-definidos, com 20% do treinamento separado para validação. A estratificação garantiu proporção similar entre classes em todos os subconjuntos.")

# 3.5
doc.add_page_break()
add_title("3.5 Modelos Implementados", level=2)
add_paragraph("Foram implementados três algoritmos de aprendizado de máquina:")

# 3.5.1
add_title("3.5.1 XGBoost", level=3)
add_paragraph("<b>Configuração ToN-IoT</b>: <i>n_estimators=200</i>, <i>max_depth=5</i>, <i>learning_rate=0.05</i>, <i>subsample=0.8</i>, <i>colsample_bytree=0.8</i>, <i>gamma=1.0</i>. <b>Configuração UNSW-NB15</b>: <i>n_estimators=300</i>, <i>max_depth=4</i>, demais parâmetros idênticos.")

# 3.5.2
add_title("3.5.2 LightGBM", level=3)
add_paragraph("<b>Configuração ToN-IoT</b>: <i>n_estimators=200</i>, <i>num_leaves=31</i>, <i>learning_rate=0.05</i>, estratégia <i>leaf-wise</i>. <b>Configuração UNSW-NB15</b>: <i>n_estimators=300</i>, demais parâmetros mantidos. Durante o treinamento, utilizou-se conjunto de validação via <i>eval_set</i> para detecção de <i>overfitting</i>.")

# 3.5.3
add_title("3.5.3 SVM", level=3)
add_paragraph("Utilizou-se <b>LinearSVC</b> com <i>dual=False</i>, <i>C=1.0</i>, <i>max_iter</i> variando entre 1000-2000. O modelo foi encapsulado em <i>CalibratedClassifierCV</i> para obtenção de probabilidades de predição.")

# 3.6
add_title("3.6 Métricas de Avaliação", level=2)
add_paragraph("Os modelos foram avaliados através de: <b>Acurácia</b> (proporção de predições corretas), <b>Precisão</b> (proporção de predições positivas verdadeiras), <b>Recall</b> (proporção de casos positivos identificados) e <b>F1-Score</b> (média harmônica entre precisão e <i>recall</i>). No contexto de segurança cibernética, o <i>Recall</i> é particularmente crítico, pois falsos negativos representam maior risco.")

# 3.7
add_title("3.7 Estratégia Experimental", level=2)
add_paragraph("A validação seguiu protocolo rigoroso: (1) Treinamento inicial com hiperparâmetros otimizados; (2) Validação intermediária; (3) Teste final em conjunto isolado; (4) Análise de <i>overfitting</i> (gap > 0,05 como alerta); (5) Exportação automática de relatórios, matrizes de confusão e gráficos de importância de <i>features</i>. Todos os experimentos utilizaram <i>random_state=42</i> para reprodutibilidade total.")

# Salvar documento
doc.save("Secao3_Materiais_e_Metodos.docx")
print("✅ Documento Word gerado: Secao3_Materiais_e_Metodos.docx")